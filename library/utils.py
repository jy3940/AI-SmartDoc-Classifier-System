"library/utils.py - File operations and analysis utilities"

from pathlib import Path
import os
from typing import Tuple
import zipfile
import xml.etree.ElementTree as ET

class FileAnalyzer:
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.xlsx', '.csv', '.json', '.xml', '.txt', '.png', '.jpg', '.jpeg'}
    
    @staticmethod
    def is_empty(file_path):
        "Check if file is empty or corrupted"
        try:
            return os.path.getsize(file_path) == 0
        except:
            return True
    
    @staticmethod
    def get_format(file_path):
        "Get file format"
        return Path(file_path).suffix.lower()
    
    @staticmethod
    def is_supported(file_path):
        "Check if format is supported"
        return FileAnalyzer.get_format(file_path) in FileAnalyzer.SUPPORTED_FORMATS


class FileOrganizer:
    "Organize files into folders by category"
    
    def __init__(self, output_dir):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _get_next_number(self, folder, stem):
        "Find the highest number used for a file with the same stem"
        max_num = 0
        for file in folder.glob(f"{stem}*"):
            # Extract number from filename
            name = file.stem
            if name == stem:
                continue  # Skip if no number
            suffix = name[len(stem):]
            try:
                num = int(suffix)
                max_num = max(max_num, num)
            except ValueError:
                pass  # Skip files that don't have pure number suffix
        return max_num
    
    def organize(self, file_path, category, mode="copy", rename_mode=False):
        "Move/copy file to category folder with optional sequential renaming"
        try:
            source = Path(file_path)
            if not source.exists():
                return None, False
            
            # Create category folder
            dest_folder = self.output_dir / category
            dest_folder.mkdir(exist_ok=True)
            
            # Handle file naming
            if rename_mode:
                # Sequential numbering: resume1, resume2, etc. (using category name as base)
                category_lower = category.lower()
                suffix = source.suffix
                
                # Find the highest number already used for this category
                max_num = self._get_next_number(dest_folder, category_lower)
                next_num = max_num + 1
                
                dest_file = dest_folder / f"{category_lower}{next_num}{suffix}"
            else:
                # Standard naming with collision handling (file_1, file_2, etc.)
                dest_file = dest_folder / source.name
                counter = 1
                while dest_file.exists():
                    stem = source.stem
                    suffix = source.suffix
                    dest_file = dest_folder / f"{stem}_{counter}{suffix}"
                    counter += 1
            
            # Copy or move
            if mode == "move":
                source.rename(dest_file)
            else:
                import shutil
                shutil.copy2(source, dest_file)
            
            return str(dest_file), True
        except Exception as e:
            return None, False


class TextExtractor:
    "Extract text from various file formats"
    
    @staticmethod
    def extract(file_path):
        "Extract text from file"
        try:
            fmt = Path(file_path).suffix.lower()
            text = ""
            
            if fmt == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            elif fmt == '.pdf':
                try:
                    from PyPDF2 import PdfReader
                    with open(file_path, 'rb') as f:
                        pdf = PdfReader(f)
                        for page in pdf.pages[:10]:  # First 10 pages
                            text += page.extract_text() or ""
                except:
                    text = ""

                if not text.strip():
                    text = TextExtractor._ocr_pdf(file_path)
            
            elif fmt == '.docx':
                text = TextExtractor._extract_docx(file_path)
            
            elif fmt in ['.xlsx', '.csv']:
                try:
                    import pandas as pd
                    if fmt == '.xlsx':
                        df = pd.read_excel(file_path, sheet_name=0)
                    else:
                        df = pd.read_csv(file_path)
                    text = " ".join(df.astype(str).values.flatten())
                except:
                    text = ""
            
            elif fmt == '.json':
                import json
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    data = json.load(f)
                    text = str(data)
            
            elif fmt == '.xml':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()

            elif fmt in ['.png', '.jpg', '.jpeg']:
                text = TextExtractor._ocr_image(file_path)
            
            return text[:5000]  # Limit to 5000 chars
        except:
            return ""

    @staticmethod
    def _extract_docx(file_path):
        "Extract DOCX text from paragraphs/tables and fallback to OOXML text nodes."
        text_parts = []
        try:
            from docx import Document

            doc = Document(file_path)

            # Body paragraphs
            for p in doc.paragraphs:
                p_text = (p.text or "").strip()
                if p_text:
                    text_parts.append(p_text)

            # Body tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        c_text = (cell.text or "").strip()
                        if c_text:
                            text_parts.append(c_text)

            # Headers/footers may contain useful text for templates.
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    p_text = (paragraph.text or "").strip()
                    if p_text:
                        text_parts.append(p_text)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            c_text = (cell.text or "").strip()
                            if c_text:
                                text_parts.append(c_text)
                for paragraph in section.footer.paragraphs:
                    p_text = (paragraph.text or "").strip()
                    if p_text:
                        text_parts.append(p_text)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            c_text = (cell.text or "").strip()
                            if c_text:
                                text_parts.append(c_text)
        except Exception:
            pass

        direct_text = " ".join(text_parts).strip()
        if len(direct_text) >= 30:
            return direct_text

        # Fallback: parse raw OOXML text nodes (captures many textbox-based templates).
        xml_text = TextExtractor._extract_docx_xml_text(file_path)
        if len(xml_text) > len(direct_text):
            return xml_text
        return direct_text

    @staticmethod
    def _extract_docx_xml_text(file_path):
        "Extract all w:t text nodes from DOCX XML parts."
        try:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            parts = []
            with zipfile.ZipFile(file_path, "r") as zf:
                target_parts = [
                    name
                    for name in zf.namelist()
                    if name.startswith("word/")
                    and name.endswith(".xml")
                    and (
                        name == "word/document.xml"
                        or name.startswith("word/header")
                        or name.startswith("word/footer")
                    )
                ]
                for xml_part in target_parts:
                    try:
                        xml_bytes = zf.read(xml_part)
                        root = ET.fromstring(xml_bytes)
                        for node in root.findall(".//w:t", ns):
                            value = (node.text or "").strip()
                            if value:
                                parts.append(value)
                    except Exception:
                        continue
            return " ".join(parts).strip()
        except Exception:
            return ""

    @staticmethod
    def _ocr_pdf(file_path, max_pages=2):
        "OCR fallback for image-based PDFs (optional dependencies)."
        try:
            from pdf2image import convert_from_path
            import pytesseract

            poppler_path = os.environ.get("POPPLER_PATH")
            pages = convert_from_path(
                file_path,
                first_page=1,
                last_page=max_pages,
                poppler_path=poppler_path
            )
            text = " ".join(pytesseract.image_to_string(page) for page in pages)
            return text
        except Exception:
            return ""

    @staticmethod
    def _ocr_image(file_path):
        "OCR for image files (optional dependencies)."
        try:
            from PIL import Image
            import pytesseract

            image = Image.open(file_path)
            return pytesseract.image_to_string(image)
        except Exception:
            return ""
